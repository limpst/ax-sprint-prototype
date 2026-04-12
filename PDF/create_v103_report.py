"""
에이톰-AX v1.0.3 기술 보고서 — Word 문서 생성
AI 데이터 파이프라인 55개 학습 데이터셋 · 사업계획서 4장 기술 근거 보완
"""
import sys, re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 페이지 여백 ─────────────────────────────────
s = doc.sections[0]
s.page_width   = Cm(21.0)
s.page_height  = Cm(29.7)
s.left_margin  = Cm(2.0)
s.right_margin = Cm(2.0)
s.top_margin   = Cm(2.0)
s.bottom_margin= Cm(1.8)

# ── 스타일 헬퍼 ─────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F,0x49,0x7D)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(10)
    return p

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x26,0x5F,0xAB)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x2E,0x74,0xB5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, bold=False, indent=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    if bold: r.bold = True
    p.paragraph_format.left_indent = Cm(indent * 0.7)
    p.paragraph_format.space_after = Pt(3)
    return p

def bullet(text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    text = re.sub(r'^[-·•]\s*', '', text)
    r = p.add_run(text)
    r.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(level * 0.7)
    p.paragraph_format.space_after = Pt(2)
    return p

def code_block(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.font.name = 'Courier New'
    r.font.color.rgb = RGBColor(0x33,0x33,0x33)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(4)
    return p

def shd_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None, header_color='1F497D'):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    hr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(8)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd_cell(c, header_color)
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri+1]
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            c.text = str(val)
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(8)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for r in tbl.rows:
                r.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return tbl

def add_pipeline_table(title, rows_data):
    """파이프라인 테이블: ID, RAW, 학습, 모델, 결과, 활용, 효과"""
    h3(title)
    headers = ["ID", "RAW 데이터", "학습 데이터", "AI 모델", "분석 결과", "현장 활용", "기대 효과"]
    tbl = doc.add_table(rows=1+len(rows_data), cols=7)
    tbl.style = 'Table Grid'
    hr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(7)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd_cell(c, '1F497D')
    for ri, row in enumerate(rows_data):
        tr = tbl.rows[ri+1]
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            c.text = str(val)
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(7)
    # 컬럼 너비
    widths = [1.2, 2.8, 2.8, 2.5, 2.8, 2.5, 2.5]
    for ci, w in enumerate(widths):
        for r in tbl.rows:
            r.cells[ci].width = Cm(w)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run("에이톰-AX")
r.bold = True; r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1F,0x49,0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AX 기반 공공임대주택 시설물 안전·유지관리 최적화\n및 지능형 행정 자동화 솔루션 상용화")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x26,0x5F,0xAB)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run("v1.0.3 기술 보고서")
r.bold = True; r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x10,0xB9,0x81)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
r = p.add_run("AI 데이터 파이프라인 55개 학습 데이터셋\n사업계획서 4장 기술 근거 보완")
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("2026년 4월 12일")
r.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run("국토교통부 AX-SPRINT (AI 응용제품 신속 상용화 지원사업)\n주관: (주)에이톰엔지니어링 | 참여: 세종대학교 × 한국화재보험협회\n위탁: (사)첨단기술안전점검협회 | 자문: 한국건설기술연구원 × 법무법인 수호")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x66,0x66,0x66)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. 개요
# ════════════════════════════════════════════════════════════
h1("1. v1.0.3 업데이트 개요")

body("본 보고서는 에이톰-AX 플랫폼 v1.0.3의 핵심 업데이트 사항인 'AI 데이터 파이프라인 55개 학습 데이터셋'에 대한 기술 보고서입니다. 사업계획서 4장 '추진 전략 및 계획'의 기술적 근거를 보완하기 위해, 플랫폼 전체 AI 서비스의 데이터 흐름을 RAW 데이터 → 학습 데이터 → AI 모델 → 분석 결과 → 현장 활용 → 기대 효과의 6단계로 체계화하였습니다.")

h2("1.1 업데이트 요약")
add_table(
    ["항목", "내용"],
    [
        ["버전", "v1.0.3 (2026-04-12)"],
        ["핵심 변경", "AI 데이터 파이프라인 55개 학습 데이터셋 추가"],
        ["시스템 변경", "PAGE 9 '🧬 AI 데이터 파이프라인' 신설, make_ai_pipeline_samples() 함수 추가"],
        ["문서 변경", "README.md 섹션 11 신설 (55건 전체 파이프라인 테이블)"],
        ["Notion 변경", "v1.0.3 신규 페이지 생성 (기존 v1.0.2 유지)"],
        ["메뉴 확장", "사이드바 8개 → 9개 페이지"],
        ["코드 규모", "app.py 약 1,700줄+ (기존 1,487줄 대비 +213줄)"],
        ["AI 모델 종류", "25종+ (비전AI·시계열·이상탐지·앙상블·NLP·3D·최적화·융합)"],
    ],
    col_widths=[4, 13]
)

h2("1.2 변경 배경")
body("사업계획서 4장 작성 시 다음 요구사항이 제기되었습니다:")
bullet("실제 데이터(RAW) → 학습 데이터 → AI 모델 → 결과 → 활용 → 효과로 이어지는 샘플 데이터 예시 최소 30개~100개")
bullet("드론·앱 이미지, Y-MaskNet, IoT·LiDAR, 민원 데이터, Digital Twin, 이상 징후, 선제 탐지, RPA 행정, 예방조치, 대시보드 분석 등 10종 데이터 유형 커버")
bullet("전체 서비스가 연계되어 설명되는 End-to-End 파이프라인 구조 필요")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 2. 시스템 아키텍처
# ════════════════════════════════════════════════════════════
h1("2. 시스템 아키텍처 변경 사항")

h2("2.1 플랫폼 구성 (v1.0.3)")
add_table(
    ["페이지", "명칭", "핵심 기능", "비고"],
    [
        ["PAGE 1", "📊 통합 대시보드", "KPI 6개, 실시간 위험 알림 6건, 차트 4개", ""],
        ["PAGE 2", "✈ 비전AI·드론 정밀진단", "Y-MaskNet 이미지 분석, 탐지 결과, 세움터 연동", ""],
        ["PAGE 3", "🤖 RPA 행정 자동화", "관리비 80%, 계약 100%, 점검 90% 자동화", ""],
        ["PAGE 4", "📋 AI 민원 관리", "7종 자동분류, 앱 접수, 처리 통계", ""],
        ["PAGE 5", "📡 IoT·디지털 트윈", "5개 센서 실시간, 2×3 서브플롯, 이상 로그", ""],
        ["PAGE 6", "🛡 AI 사전 예방 정비", "XAI 위험도 스코어링, SHAP, 정비 지시서", "핵심 예방 엔진"],
        ["PAGE 7", "🏅 클린하우스 마일리지", "입주민 참여 인센티브, 4등급 체계", ""],
        ["PAGE 8", "🔮 Vision 2030 예측분석", "고장 예측, 에너지 최적화, 사업화 KPI", ""],
        ["PAGE 9", "🧬 AI 데이터 파이프라인", "55개 학습 데이터셋, 9개 카테고리, 25종+ AI 모델", "v1.0.3 신규"],
    ],
    col_widths=[2, 4.5, 6.5, 3]
)

h2("2.2 데이터 생성 함수 (v1.0.3)")
add_table(
    ["함수명", "행 수", "역할", "비고"],
    [
        ["make_iot_data()", "720행", "IoT 센서 시계열 (온도·습도·전력·진동·CO)", ""],
        ["make_complaints()", "80행", "민원 데이터 (7종 유형, 6단계 상태)", ""],
        ["make_inspection()", "84행", "점검 이력 (노후도 공식 적용)", ""],
        ["make_facilities()", "7행", "시설물 대장 (세움터·FMS 연동)", ""],
        ["make_damage_detections()", "8행", "AI 손상 탐지 결과 (문서 기반 하드코딩)", ""],
        ["make_billing_data()", "504행", "관리비 고지서 (RPA 자동화 대상)", ""],
        ["make_contract_data()", "40행", "계약 만료 (D-Day 자동 계산)", ""],
        ["make_milage_data()", "50행", "클린하우스 마일리지 (4등급)", ""],
        ["make_failure_forecast()", "48행", "고장 예측 (Vision 2030 ML)", ""],
        ["make_ai_pipeline_samples()", "55행", "AI 데이터 파이프라인 (9개 카테고리)", "v1.0.3 신규"],
    ],
    col_widths=[5, 1.8, 6, 2.5]
)

h2("2.3 6단계 데이터 파이프라인 구조")
code_block("① RAW 데이터 → ② 학습 데이터 → ③ AI 모델 → ④ 분석 결과 → ⑤ 현장 활용 → ⑥ 기대 효과")
code_block("         ↑                                                                    │")
code_block("         └──────────────── MLOps 피드백 루프 (결과 → 재학습) ──────────────────┘")

body("모든 파이프라인은 End-to-End로 연계되어, RAW 데이터 수집부터 현장 조치까지 자동으로 흐르며, 결과가 다시 학습 데이터로 환류되는 MLOps 구조를 갖추고 있습니다.")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. 55개 AI 데이터 파이프라인 상세
# ════════════════════════════════════════════════════════════
h1("3. 55개 AI 데이터 파이프라인 상세")

body("9개 카테고리, 55건의 파이프라인 데이터셋을 다음과 같이 정의합니다. 각 파이프라인은 RAW 데이터 수집 → 학습 데이터 전처리 → AI 모델 적용 → 분석 결과 도출 → 현장 활용 → 기대 효과의 6단계 구조를 따릅니다.")

add_table(
    ["카테고리", "건수", "주요 AI 모델", "핵심 데이터 소스"],
    [
        ["🖼 비전AI — 드론·앱 이미지", "10건", "Y-MaskNet, Thermal-MaskNet, LSTM", "드론 4K, 앱 사진, 열화상"],
        ["📡 IoT 센서·LiDAR", "8건", "AutoEncoder, LSTM, PointNet++", "온도·CO·진동·습도·전력·3D점군"],
        ["📋 민원 데이터", "6건", "KoBERT, XGBoost, EfficientNet", "민원 텍스트·사진·이력"],
        ["🏗 Digital Twin", "6건", "PINN, Antigravity, Cox, FDS", "BIM·FEM·KALIS·GPR"],
        ["🚨 이상 징후 감지", "5건", "Attention-LSTM, Isolation Forest", "복합 센서 융합 실시간"],
        ["🛡 선제 탐지", "5건", "LSTM, PCA, Seasonal ARIMA", "시계열 추적·예보 연동"],
        ["🤖 RPA 행정 자동화", "7건", "Rule Engine, LLM+RAG, RL", "관리비·계약·점검·법령"],
        ["🔰 예방조치", "5건", "XAI 스코어링, STL-XGBoost", "XAI 위험도·GPR·InSAR"],
        ["📊 대시보드 분석", "3건", "Ensemble Forecasting, Attribution", "KPI·예측·비용 비교"],
    ],
    col_widths=[4, 1.5, 5, 5.5]
)

# ── 3.1 비전 AI ──
add_pipeline_table("3.1 비전 AI — 드론·앱 시설물 손상 이미지 (10건)", [
    ["P-001", "드론 촬영 외벽 4K (강남A동 3층)", "512×512 패치 + 균열 마스크 (COCO)", "Y-MaskNet (F1=0.97)", "수직균열 0.35mm, 94%", "48h 긴급 보수 지시", "구조사고 차단, 보수비 70%↓"],
    ["P-002", "앱 업로드 지하층 천장 (마포B동)", "습윤 영역 세그멘테이션 마스크", "Y-MaskNet (IoU=0.91)", "누수 312cm², 88%", "3일 내 배관팀 배정", "하자보수 45%↓"],
    ["P-003", "드론 스캔 외벽+열화상 (서초E동)", "드라이비트 BBox + KFS 매핑", "Y-MaskNet+Thermal", "화재취약 850cm, 91%", "30일 불연재 교체", "화재 차단, 보험료 20%↓"],
    ["P-004", "드론 옥상 방수층 (노원C동)", "박리 폴리곤 + 깊이 회귀", "Multi-task Y-MaskNet", "박리 2.1mm, 89%", "방수 재시공 지시", "실내 피해 예방"],
    ["P-005", "앱 업로드 계단실 (강남A동)", "미세균열 + Antigravity 보정", "Y-MaskNet+Antigravity", "0.08mm (기준미만), 79%", "정기점검 유지", "오탐 0건"],
    ["P-006", "드론 지하주차장 (강서D동)", "누수 5종 분류 라벨", "5-class (Acc=0.93)", "배관 누수 0.45mm, 86%", "배관 교체 지시", "유지비 35%↓"],
    ["P-007", "앱 외벽 표면 (마포B동 5층)", "박락 마스크+깊이·면적", "mAP=0.89", "박락 0.55mm, 83%", "30일 부분 보수", "낙하물 사고 예방"],
    ["P-008", "드론 파노라마 (구미A동)", "대면적 세그멘테이션+두께", "Y-MaskNet+Panorama", "드라이비트 1200cm, 93%", "불연재 교체 발주", "대규모 화재 차단"],
    ["P-009", "드론 야간 열화상 (서초E동)", "열분포 히트맵+단열등급", "Thermal-MaskNet", "단열불량 3개소, 87%", "단열 보강 계획", "난방비 12%↓"],
    ["P-010", "앱 3개월 추적 (강남A동)", "균열폭 시계열 변화", "Y-MaskNet+LSTM", "성장률 0.07mm/월, 92%", "구조 긴급 평가", "위험 前 선제 조치"],
])

# ── 3.2 IoT ──
add_pipeline_table("3.2 IoT 센서·LiDAR 데이터 (8건)", [
    ["P-011", "온도센서 30일 (마포B동)", "정상/이상 라벨", "AutoEncoder", "32.1°C 이상 감지", "환기 권고+HVAC 점검", "열사병 차단, 에너지 8%↑"],
    ["P-012", "CO 농도 실시간 (마포B동)", "급등 패턴 라벨", "LSTM+Anomaly", "16.2ppm (62% 초과)", "환기팀 출동+대피", "CO 중독 차단"],
    ["P-013", "진동 3축 100Hz (강남A동)", "주파수+위험등급", "Random Forest", "0.85g (70% 초과)", "구조팀 긴급 점검", "붕괴 위험 차단"],
    ["P-014", "습도 15분 간격 (노원C동)", "습도-누수 상관", "Gradient Boosting", "78.4%, 누수확률 73%", "3일 내 점검 예약", "수리비 50%↓"],
    ["P-015", "LiDAR 3D 1억점 (강남A동)", "3D 메시+균열 라벨", "PointNet++ (mIoU=0.88)", "변형 3개소, 12mm", "Digital Twin 갱신", "탐지율 15%↑"],
    ["P-016", "전력 스마트미터 (전 단지)", "패턴+이상 소비", "K-Means+Isolation", "5.8kWh (240%)", "전기안전 점검", "전기 화재 예방"],
    ["P-017", "LiDAR SLAM (노원C동 옥상)", "시계열 변위맵", "ICP+Change Det.", "방수층 부풀음 +18%", "우기 전 긴급 보수", "옥상 누수 차단"],
    ["P-018", "5개 센서 융합 72h", "복합 이상 시나리오", "Multivariate LSTM", "화재 전조 패턴", "소방+대피+스프링클러", "대응 70% 단축"],
])

# ── 3.3 민원 ──
add_pipeline_table("3.3 공공임대 민원 데이터 (6건)", [
    ["P-019", "민원 텍스트 (물방울)", "유형+위험도+위치 NER", "KoBERT (Acc=0.92)", "누수/중위험, 0.89", "배관팀 배정+3일 방문", "처리시간 67%↓"],
    ["P-020", "텍스트 (외벽 떨어짐)", "긴급 키워드 가중치", "KoBERT+Rule", "균열/고위험, 0.94", "즉시 전문가 배정", "고위험 누락 0건"],
    ["P-021", "6개월 이력 480건", "빈도+노후도 피처", "XGBoost (RMSE=2.1)", "다음달 누수 12건", "선제 점검+자재 확보", "민원 30%↓"],
    ["P-022", "감성 분석 데이터", "텍스트+만족도 매핑", "KoBERT (F1=0.87)", "부정→긍정 78%", "부정 민원 우선 처리", "분쟁 40%↓"],
    ["P-023", "민원 첨부 사진", "5종 손상유형 매핑", "EfficientNet (0.86)", "곰팡이 12cm²", "환경팀+환기 안내", "수동 판독 80%↓"],
    ["P-024", "처리 이력 타임스탬프", "단계별 소요시간", "LightGBM (MAE=0.3일)", "예상 2.1일 (지연)", "추가 인력+지연 안내", "SLA 95%→98%"],
])

# ── 3.4 Digital Twin ──
add_pipeline_table("3.4 Digital Twin 데이터 (6건)", [
    ["P-025", "BIM+LiDAR 현황", "설계 vs 현황 차분", "3D Point Cloud Match", "기울기 0.15°, 8mm", "DT 갱신+구조 경보", "사각지대 커버"],
    ["P-026", "세종대 비선형 FEM", "FEM+실측 보정", "PINN (대리모델)", "10분 내 시뮬 (기존 8h)", "실시간 안전 판정", "해석 98% 단축"],
    ["P-027", "에너지 12개월+기상", "소비+외기온·일사량", "GBR (R²=0.94)", "892→743만 (-17%)", "HVAC 자동 최적화", "에너지 15%↓"],
    ["P-028", "KALIS-FMS 30년", "노후도 곡선+이력", "Antigravity 3중검증", "44.6→5년후 35.2점", "장기 보수+예산 확보", "수명 10년 연장"],
    ["P-029", "배관 도면+GPR+누수", "재질·경과·토질별", "Cox 비례위험", "3년 내 파손 42%", "배관 교체 우선순위", "긴급 파열 방지"],
    ["P-030", "외벽 3D+화재 시뮬", "재질·면적별 확산", "FDS+ML 대리모델", "3분 내 3개층 확산", "소방 시뮬+대피 훈련", "대응 50% 단축"],
])

# ── 3.5 이상 징후 ──
add_pipeline_table("3.5 실시간 이상 징후 감지 (5건)", [
    ["P-031", "5개 센서 동시 이상", "복합 시나리오 패턴", "Attention-LSTM", "화재 전조 0.91", "소방+대피+스프링클러", "차단율 94%"],
    ["P-032", "인접 공사 진동 72h", "크기·주파수별 영향", "FFT+Random Forest", "에너지 초과, 주의", "저감 요청+정밀 점검", "진동 손상 차단"],
    ["P-033", "전력 급변 (심야)", "시간대별+이상 라벨", "Isolation+DBSCAN", "5.8kWh (276%)", "전기점검+누전 모니터", "전기 화재 예방"],
    ["P-034", "가스 감지 (0.1초)", "급등+3종 원인", "1D-CNN (F1=0.91)", "미세 누출 0.3%", "사전 경고+48h 점검", "가스 폭발 차단"],
    ["P-035", "집수정 수위 5분", "수위+강우 결합", "ARIMA+ML", "6h 후 경고 예측", "펌프 가동+차량 이동", "침수 차단"],
])

# ── 3.6 선제 탐지 ──
add_pipeline_table("3.6 선제 탐지 예시 데이터 (5건)", [
    ["P-036", "균열 6개월 추적", "폭 시계열 변화", "LSTM 성장 예측", "3개월 후 0.5mm", "기준 前 선제 보수", "보수비 70%↓"],
    ["P-037", "배관 유량·압력 30일", "상관 변화 패턴", "PCA+Hotelling T²", "누수 전조 78%", "내시경+자재 확보", "파열 전 교체"],
    ["P-038", "계절별 열화상 4회", "동결-해동+균열", "Seasonal ARIMA", "봄 균열 87% 예측", "동절기 후 집중 점검", "긴급 보수 50%↓"],
    ["P-039", "승강기 로그+진동", "고장 전 7일 패턴", "XGBoost 7일 선행", "브레이크 수명 14일", "점검 앞당김+부품", "정지 사고 0건"],
    ["P-040", "옥상 수분+강우예보", "침투+강우 상관", "Weather-Aware ML", "방수 실패 65%", "우기 전 보수", "누수 5→0건"],
])

# ── 3.7 RPA ──
add_pipeline_table("3.7 RPA 행정 자동화 데이터 (7건)", [
    ["P-041", "504세대 관리비 원장", "청구 규칙+이상 패턴", "Rule+Anomaly", "504건 (12건 보정)", "80% 자동 발행", "8h→1.5h (81%↓)"],
    ["P-042", "40건 계약 만료", "갱신 확률 피처", "Logistic (AUC=0.89)", "퇴거 3건 식별", "100% 알림+대기자", "누락 0건"],
    ["P-043", "7단지 시설물 대장", "노후도→점검주기", "규칙+ML 최적화", "월간 일정 자동", "90% 자동+앱 할당", "점검 50%↓"],
    ["P-044", "민원 접수 데이터", "담당자 전문분야", "Hungarian+ML", "최적 담당자 선정", "75% 자동 배정", "2h→5min"],
    ["P-045", "점검 결과+법령", "템플릿+RAG 인덱스", "LLM+RAG", "보고서 초안 (80%↓)", "수호 포맷→공문서", "작성비 80%↓"],
    ["P-046", "연체 현황 24개월", "연체+납부 확률", "GBM (AUC=0.85)", "연체 예상 23건", "사전 안내+분납", "연체 15%→8%"],
    ["P-047", "공용부 전력", "시간대별+절전 구간", "RL HVAC 제어", "야간 50%↓+최적화", "BMS 자동 연동", "1,200만원/년↓"],
])

# ── 3.8 예방조치 ──
add_pipeline_table("3.8 예방조치 데이터 (5건)", [
    ["P-048", "XAI 위험도 88점", "7개 지표+SHAP", "XAI 스코어링 v2.0", "경과 30%+IoT 20%", "48h 긴급 정비 지시", "차단율 94%"],
    ["P-049", "드라이비트+KFS", "재질·면적→화재 등급", "Multi-criteria", "서초E A등급", "불연재+소화설비", "화재 차단"],
    ["P-050", "누수+배관+IoT습도", "복합→재발 확률", "Ensemble Stacking", "재발 73% (3개월)", "배관 전면 교체", "반복 누수 차단"],
    ["P-051", "GPR+InSAR+지하수", "다변량+침하 이벤트", "STL-XGBoost", "침하 확률 0.68", "GPR 탐사+도로 제한", "지반침하 차단"],
    ["P-052", "7단지 위험도 종합", "추이+보수 효과", "Portfolio Risk", "고3·중2·정상2", "예산 배분+경영 보고", "ROI 340%"],
])

# ── 3.9 대시보드 ──
add_pipeline_table("3.9 대시보드 분석 데이터 (3건)", [
    ["P-053", "전체 운영 데이터", "KPI 집계+목표 달성", "실시간+Trend", "점검 94.2%, 1.2일", "통합 대시보드+보고", "보고 90%↓"],
    ["P-054", "Vision 2030 예측", "과거+예측 시계열", "Ensemble Forecast", "고장 30%↓, 에너지 15%↓", "전략+예산 배분", "선제 투자 지원"],
    ["P-055", "AI 전·후 비용 비교", "절감+기여도 분해", "Cost-Benefit", "연간 7,350만원 절감", "ROI 보고서+확산", "ROI 340%, 8개월"],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. 적용 AI 모델 종합 현황
# ════════════════════════════════════════════════════════════
h1("4. 적용 AI 모델 종합 현황 (25종+)")

add_table(
    ["카테고리", "주요 모델", "적용 영역", "핵심 성능"],
    [
        ["비전 AI", "Y-MaskNet, EfficientNet, PointNet++, Thermal-MaskNet", "균열·누수·드라이비트·박락·단열불량", "F1=0.97, mAP≥0.92"],
        ["시계열 예측", "LSTM, ARIMA, Seasonal ARIMA, Cox 모델", "균열 성장, 센서 이상, 침하, 배관 수명", "RMSE=2.1, MAE=0.3일"],
        ["이상 탐지", "AutoEncoder, Isolation Forest, PCA, 1D-CNN", "IoT 센서, 전력·가스·수위", "FP ≤5%"],
        ["앙상블 학습", "XGBoost, LightGBM, Random Forest, Stacking", "위험도·민원·연체·고장 예측", "AUC≥0.85"],
        ["자연어 처리", "KoBERT, LLM+RAG", "민원 분류, 감성분석, 보고서 생성", "Acc=0.92, 80% 단축"],
        ["3D/공간", "PointNet++, ICP, FDS, PINN", "Digital Twin, 구조 해석, 화재 시뮬", "mIoU=0.88, 98% 단축"],
        ["최적화", "RL, Hungarian, Portfolio Risk", "HVAC 제어, 인력 배정, 예산 배분", "ROI 340%"],
        ["융합 예측", "STL-XGBoost, Multimodal LSTM", "지반침하, 복합 센서 이상", "차단율 94%"],
    ],
    col_widths=[2.5, 5, 4.5, 4]
)

h1("5. 핵심 원칙 및 품질 보증")

h2("5.1 데이터 파이프라인 4대 원칙")
bullet("End-to-End 연계 — 모든 RAW 데이터가 최종 현장 조치까지 자동으로 흐름")
bullet("XAI 투명성 — AI 판정 근거(SHAP)를 항상 함께 제공하여 현장 신뢰 확보")
bullet("Antigravity 보정 — KALIS-FMS 30년 이력 + 일송 설계 + 세종대 FEM 3중 검증 → 오탐 0건")
bullet("실시간 피드백 — 결과가 다시 학습 데이터로 환류 → 모델 지속 개선 (MLOps)")

h2("5.2 Antigravity 엔진 3중 교차검증")
add_table(
    ["검증 레이어", "데이터 소스", "기여도", "역할"],
    [
        ["① 이력 기반", "KALIS-FMS 20~30년 결함 이력", "40%", "노후화 패턴 학습, 결함 전이 예측"],
        ["② 설계 기반", "일송건축 설계도면 DB", "35%", "설계 의도 대비 현황 편차 분석"],
        ["③ 구조해석", "세종대 비선형 FEM", "25%", "물리 법칙 기반 구조 안전성 검증"],
    ],
    col_widths=[3, 5, 2, 6]
)

h2("5.3 성능 지표 (TRL 8 실증 기준)")
add_table(
    ["지표", "목표", "현재 달성", "비고"],
    [
        ["결함 탐지율", "≥95%", "93.1%", "Y-MaskNet + Antigravity"],
        ["False Positive", "0건", "0건", "3중 교차검증 효과"],
        ["균열 해상도", "0.2mm", "0.2mm", "드론 4K + AI 세그멘테이션"],
        ["민원처리 시간", "1일", "1.2일", "기존 3일 → 67% 단축"],
        ["RPA 자동화율", "80%", "81.5%", "관리비·계약·점검 통합"],
        ["보고서 생성", "36분/건", "36분/건", "기존 3시간 → 80% 단축"],
        ["사고 전환 차단율", "≥90%", "94%", "XAI 조기 경보 효과"],
        ["예상 ROI", "300%", "340%", "투자 회수 기간 8개월"],
    ],
    col_widths=[4, 3, 3, 6]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. 버전 히스토리
# ════════════════════════════════════════════════════════════
h1("6. 버전 히스토리")

add_table(
    ["버전", "일자", "핵심 변경 사항"],
    [
        ["v1.0.0", "2026-04-07", "초기 프로토타입 (8페이지, 9개 데이터함수, 약 1,228줄)"],
        ["v1.0.1", "2026-04-11", "손해사정 개념 삭제, 법무법인 수호·이화영 박사·LLM+RAG·KALIS-FMS·세움터 추가"],
        ["v1.0.2", "2026-04-11", "Antigravity 기술적 해자, 정책 기여도, 연차별 KPI, KALIS-FMS 노후화 곡선 추가"],
        ["v1.0.3", "2026-04-12", "AI 데이터 파이프라인 55개 학습 데이터셋, PAGE 9 신설, 25종+ AI 모델 명세"],
    ],
    col_widths=[2, 3, 11]
)

h1("7. 기대효과 종합")

add_table(
    ["분야", "AS-IS", "TO-BE (AI 적용)", "개선율"],
    [
        ["민원 처리 시간", "3일", "1일", "67% 단축"],
        ["점검 업무 시간", "8시간/회", "4시간/회", "50% 단축"],
        ["사후 보수 비용", "기준", "△20%", "20% 절감"],
        ["에너지 비용", "기준", "△15%", "15% 절감"],
        ["관리비 발행", "8시간/회", "1.5시간/회", "81% 단축"],
        ["보고서 작성", "3시간/건", "36분/건", "80% 단축"],
        ["사고 전환 차단", "사후 대응", "사전 AI 차단 94%", "사전 예방 전환"],
        ["예상 연간 절감", "-", "7,350만원", "ROI 340%"],
    ],
    col_widths=[3.5, 3, 4, 3]
)

body("")
body("본 보고서는 에이톰-AX 플랫폼 v1.0.3의 기술 변경 사항을 정리한 것으로, 국토교통부 AX-SPRINT 사업계획서 4장 '추진 전략 및 계획'의 기술적 근거 자료로 활용됩니다.")
body("")
body("— 끝 —", bold=True)

# ════════════════════════════════════════════════════════════
# 저장
# ════════════════════════════════════════════════════════════
out = "에이톰-AX_v1.0.3_기술보고서.docx"
doc.save(out)
print(f"✅ 기술 보고서 생성 완료: {out}")
print(f"   위치: PDF/{out}")
